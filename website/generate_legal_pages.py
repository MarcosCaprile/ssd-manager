"""Generate publishable static legal pages from the reviewed Word documents."""

from html import escape
from pathlib import Path
import re
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parent
DOCS = ROOT.parent / "docs" / "legal-release-package" / "word"
DOWNLOADS = ROOT / "downloads"
PRODUCT_ORIGIN = "https://ssd-manager.minutmate.com"
LINK_TOKEN = re.compile(
    r"https://[^\s<>]+|[A-Za-z0-9.!#$%&'*+/=?^_`{|}~-]+@[A-Za-z0-9](?:[A-Za-z0-9-]{0,61}[A-Za-z0-9])?(?:\.[A-Za-z0-9](?:[A-Za-z0-9-]{0,61}[A-Za-z0-9])?)+"
)

PAGES = {
    "impressum": ("01_SSD_Manager_Impressum.docx", "Impressum"),
    "datenschutz": ("02_SSD_Manager_Datenschutzerklaerung.docx", "Datenschutzerklärung"),
    "nutzungsbedingungen": ("03_SSD_Manager_Nutzungsbedingungen.docx", "Nutzungsbedingungen"),
    "avv": ("05_SSD_Manager_AVV_Schule.docx", "Auftragsverarbeitungsvertrag"),
    "toms": ("06_SSD_Manager_TOMs.docx", "Technische und organisatorische Maßnahmen"),
    "unterauftragnehmer": ("07_SSD_Manager_Unterauftragnehmerliste.docx", "Unterauftragnehmer"),
    "loeschkonzept": ("08_SSD_Manager_Loesch_und_Aufbewahrungskonzept.docx", "Lösch- und Aufbewahrungskonzept"),
}


def linkify(value: str) -> str:
    """Escape text and turn each URL or email address into exactly one link."""
    parts = []
    position = 0
    for match in LINK_TOKEN.finditer(value):
        parts.append(escape(value[position:match.start()]))
        token = match.group(0)
        bare = token.rstrip(".,;:·")
        suffix = token[len(bare):]
        if bare.startswith(PRODUCT_ORIGIN):
            href = bare[len(PRODUCT_ORIGIN):] or "/"
        elif bare.startswith("https://"):
            href = bare
        else:
            href = f"mailto:{bare}"
        parts.append(f'<a href="{escape(href, quote=True)}">{escape(bare)}</a>{escape(suffix)}')
        position = match.end()
    parts.append(escape(value[position:]))
    return "".join(parts).replace("\n", "<br>")


def is_internal(paragraph: str) -> bool:
    return paragraph.startswith("Prüfhinweis:") or "vor Einsatz rechtlich und schulrechtlich freigeben" in paragraph


def document_body(path: Path) -> str:
    doc = Document(path)
    parts = []
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if not text or is_internal(text):
            continue
        style = paragraph.style.name if paragraph.style else ""
        if index == 0:
            continue
        if style.startswith("Heading"):
            level = 2 if style == "Heading 1" else 3
            parts.append(f"<h{level}>{linkify(text)}</h{level}>")
        elif style.startswith("List Bullet"):
            parts.append(f'<p class="list-item">{linkify(text)}</p>')
        else:
            parts.append(f"<p>{linkify(text)}</p>")
    for table in doc.tables:
        rows = []
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(is_internal(value) for value in values):
                continue
            if values == ["Produkt", "SSD Manager"] or values[0] in {"Stand", "Status"}:
                continue
            cells = "".join(f"<td>{linkify(value)}</td>" for value in values)
            rows.append(f"<tr>{cells}</tr>")
        if rows:
            parts.append('<div class="table-wrap"><table>' + "".join(rows) + "</table></div>")
    return "\n".join(parts)


def shell(title: str, body: str, download: str) -> str:
    return f'''<!doctype html>
<html lang="de">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{escape(title)} · SSD Manager</title>
  <meta name="description" content="{escape(title)} für SSD Manager von MinutMate.">
  <link rel="icon" href="/assets/app-icon.png">
  <link rel="stylesheet" href="/styles.css">
</head>
<body>
  <header class="site-header"><a class="brand" href="/"><img src="/assets/app-icon.png" alt=""><span>SSD Manager</span></a><nav aria-label="Hauptnavigation"><a href="/#funktionen">Funktionen</a><a href="/#schulen">Für Schulen</a><a href="/support/">Support</a></nav></header>
  <main class="legal-page">
    <div class="eyebrow">Rechtliches</div>
    <h1>{escape(title)}</h1>
    <p class="legal-intro">Stand: 26. Juli 2026 · Die schulbezogenen Angaben werden bei Einrichtung der jeweiligen Schulumgebung ergänzt.</p>
    <article class="legal-card">{body}</article>
    <a class="download-link" href="/downloads/{escape(download)}" download>Bearbeitbare Word-Fassung herunterladen</a>
  </main>
  <footer><div><strong>SSD Manager</strong><br>Ein Produkt von MinutMate</div><div class="footer-links"><a href="/impressum/">Impressum</a><a href="/datenschutz/">Datenschutz</a><a href="/nutzungsbedingungen/">Nutzungsbedingungen</a><a href="/konto-loeschen/">Kontolöschung</a><a href="/avv/">AVV</a><a href="/toms/">TOMs</a><a href="/unterauftragnehmer/">Unterauftragnehmer</a><a href="/loeschkonzept/">Löschkonzept</a></div></footer>
</body>
</html>'''


def main():
    DOWNLOADS.mkdir(exist_ok=True)
    for slug, (filename, title) in PAGES.items():
        source = DOCS / filename
        output = ROOT / slug
        output.mkdir(exist_ok=True)
        (output / "index.html").write_text(shell(title, document_body(source), filename), encoding="utf-8")
        shutil.copy2(source, DOWNLOADS / filename)
        print(f"Generated /{slug}/ from {filename}")


if __name__ == "__main__":
    main()
