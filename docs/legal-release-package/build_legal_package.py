from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "word"
OUT.mkdir(parents=True, exist_ok=True)

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 98, 108)
LIGHT = "F2F4F7"
TODAY = "26.07.2026"
OPERATOR = "Marcos Caprile Santos, handelnd unter MinutMate"
ADDRESS = "Heltorfer Mark 131, 40489 Düsseldorf, Deutschland"


def shade(cell, fill=LIGHT):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120, repeat_header=True):
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    for tag, value in [("w:tblW", total), ("w:tblInd", indent)]:
        node = tbl_pr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
    if repeat_header and table.rows:
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)


def configure(doc, short_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    hp = section.header.paragraphs[0]
    hp.text = f"SSD Manager | {short_title}"
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in hp.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
    fp = section.footer.paragraphs[0]
    fp.text = "MinutMate · Arbeits- und Prüffassung · vertraulich"
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in fp.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.color.rgb = GRAY


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = DARK
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    r2 = p2.add_run(subtitle)
    r2.font.name = "Calibri"
    r2.font.size = Pt(13)
    r2.font.color.rgb = GRAY
    meta = doc.add_table(rows=3, cols=2)
    meta.autofit = False
    set_table_geometry(meta, [1800, 7560], repeat_header=False)
    rows = [("Produkt", "SSD Manager"), ("Stand", TODAY), ("Status", "Prüffassung – vor Einsatz rechtlich und schulrechtlich freigeben")]
    for i, (label, value) in enumerate(rows):
        set_cell_width(meta.cell(i, 0), 1800)
        set_cell_width(meta.cell(i, 1), 7560)
        meta.cell(i, 0).text = label
        meta.cell(i, 1).text = value
        shade(meta.cell(i, 0))
        meta.cell(i, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        meta.cell(i, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        meta.cell(i, 0).paragraphs[0].runs[0].bold = True
    doc.add_paragraph()


def add_note(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_cell_width(table.cell(0, 0), 9360)
    shade(table.cell(0, 0), "FFF4CE")
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run("Prüfhinweis: ")
    r.bold = True
    p.add_run(text)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(.5)
        p.paragraph_format.first_line_indent = Inches(-.25)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(.5)
        p.paragraph_format.first_line_indent = Inches(-.25)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.167
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    widths = widths or [9360 // len(headers)] * len(headers)
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        set_cell_width(table.cell(0, i), widths[i])
        table.cell(0, i).text = header
        shade(table.cell(0, i))
        table.cell(0, i).paragraphs[0].runs[0].bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_width(cells[i], widths[i])
            cells[i].text = str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def signatures(doc):
    doc.add_heading("Unterschriften", level=1)
    add_table(doc, ["Schule / Schulträger", "MinutMate"], [
        ("Ort, Datum: ______________________________\nName/Funktion: ___________________________\nUnterschrift: ____________________________",
         "Ort, Datum: ______________________________\nMarcos Caprile Santos, Inhaber\nUnterschrift: ____________________________")
    ], [4680, 4680])


def save(filename, title, subtitle, sections, note=True, sign=False):
    doc = Document()
    configure(doc, title)
    add_title(doc, title, subtitle)
    if note:
        add_note(doc, "Dieses Dokument ist eine belastbare Arbeits- und Prüffassung, ersetzt aber keine individuelle Rechtsberatung. Schule/Schulträger, zuständiger Datenschutzbeauftragter und eine qualifizierte Rechtsberatung müssen die endgültige Fassung vor Verarbeitung echter Schuldaten prüfen.")
    for heading, body in sections:
        doc.add_heading(heading, level=1)
        for block in body:
            kind = block[0]
            if kind == "p":
                doc.add_paragraph(block[1])
            elif kind == "h2":
                doc.add_heading(block[1], level=2)
            elif kind == "bullets":
                add_bullets(doc, block[1])
            elif kind == "numbered":
                add_numbered(doc, block[1])
            elif kind == "table":
                add_table(doc, block[1], block[2], block[3] if len(block) > 3 else None)
    if sign:
        signatures(doc)
    doc.core_properties.title = title
    doc.core_properties.subject = "SSD Manager Datenschutz- und Vertragsunterlagen"
    doc.core_properties.author = "MinutMate"
    doc.core_properties.keywords = "SSD Manager, DSGVO, Schule, Prüffassung"
    doc.save(OUT / filename)


documents = []


def register(filename, title, purpose, status):
    documents.append((filename, title, purpose, status))


save("01_SSD_Manager_Impressum.docx", "Impressum SSD Manager", "Anbieterkennzeichnung nach § 5 DDG", [
    ("Angaben zum Diensteanbieter", [("p", f"{OPERATOR}\nEinzelunternehmen\n{ADDRESS}")]),
    ("Kontakt", [("p", "Telefon: +49 172 7723349\nProduktkontakt: info@ssd-manager.minutmate.com\nProduktsupport: support@ssd-manager.minutmate.com\nAllgemeiner MinutMate-Kontakt: info@minutmate.com\nVerwaltung: verwaltung@minutmate.com\nDatenschutzkontakt: datenschutz@minutmate.com")]),
    ("Unternehmens- und Steuerangaben", [("p", "MinutMate wird als Einzelunternehmen geführt. Eine Handelsregistereintragung besteht derzeit nicht. Umsatzsteuer-Identifikationsnummer gemäß § 27a UStG: DE462926396.")]),
    ("Verbraucherstreitbeilegung", [("p", "MinutMate ist nicht verpflichtet und nicht bereit, an Streitbeilegungsverfahren vor einer Verbraucherschlichtungsstelle teilzunehmen.")]),
    ("Geltungsbereich und Links", [("p", "Dieses Impressum gilt für die SSD-Manager-App, ihre produktbezogenen Webseiten und Subdomains. Vor Veröffentlichung sind die endgültigen URLs für Datenschutzerklärung, Nutzungsbedingungen und Support einzutragen.")]),
], note=True)
register("01_SSD_Manager_Impressum.docx", "Impressum", "Öffentliche Anbieterkennzeichnung", "URL ergänzen und anwaltlich prüfen")

save("02_SSD_Manager_Datenschutzerklaerung.docx", "Datenschutzerklärung SSD Manager", "Informationen nach Art. 13 und 14 DSGVO", [
    ("1. Verantwortlichkeit", [("p", f"Für eigene Verarbeitungen von MinutMate verantwortlich: {OPERATOR}, {ADDRESS}. Datenschutzkontakt: datenschutz@minutmate.com. Für die schulische Organisation des Schulsanitätsdienstes, die Auswahl der Nutzenden, Rollen, Dienstplanung und Kommunikationsinhalte ist regelmäßig die jeweilige Schule oder ihr Schulträger Verantwortlicher; MinutMate verarbeitet diese Daten insoweit als Auftragsverarbeiter."), ("p", "Die Kontaktdaten der jeweils verantwortlichen Schule werden bei der schulischen Bereitstellung mitgeteilt. Der Datenschutzkontakt von MinutMate ist keine Bezeichnung eines gesetzlich bestellten Datenschutzbeauftragten.")]),
    ("2. Datenkategorien", [("table", ["Kategorie", "Beispiele"], [
        ("Konto- und Stammdaten", "Name, Benutzername, schulische E-Mail, Schule, Rolle, Status, Qualifikationsbeginn"),
        ("Dienst- und Organisationsdaten", "Diensttage, Zuweisungen, Krankmeldungen, Status und Historie"),
        ("Kommunikations- und Moderationsdaten", "Schulweite Ankündigungstexte, Absenderzuordnung, Fotos, Dateien, Meldegrund, optionale Zusatzangaben und Bearbeitungsstatus"),
        ("Geräte- und Sicherheitsdaten", "Installationskennung, Geräteinformationen, Sitzungshashes, Push-Token, Login- und Auditdaten"),
        ("Support- und Vertragsdaten", "Anfragen, Ansprechpartner, Vertrags-, Abrechnungs- und Nachweisdaten"),
    ], [2500, 6860])]),
    ("3. Zwecke und Rechtsgrundlagen", [("p", "Schulische Verarbeitungen erfolgen auf der von Schule/Schulträger festgelegten Rechtsgrundlage, insbesondere der jeweiligen schul- und datenschutzrechtlichen Aufgabennorm. MinutMate entscheidet diese Rechtsgrundlage nicht."), ("bullets", [
        "Bereitstellung und Administration der Schulumgebung, Rollen und Konten.",
        "Planung und Nachweis von Diensten des Schulsanitätsdienstes.",
        "Schulweite organisatorische Kommunikation einschließlich angeforderter Anhänge.",
        "Authentifizierung, IT-Sicherheit, Missbrauchsabwehr, Fehleranalyse und Support.",
        "Vertragsverwaltung, Abrechnung und Erfüllung gesetzlicher Pflichten in eigener Verantwortung von MinutMate.",
    ]), ("p", "Soweit MinutMate eigene Verarbeitungen vornimmt, kommen insbesondere Art. 6 Abs. 1 Buchst. b, c und f DSGVO in Betracht. Die konkrete Interessenabwägung ist im Verzeichnis der Verarbeitungstätigkeiten zu dokumentieren.")]),
    ("4. Empfänger und Dienstleister", [("bullets", [
        "Railway Corporation: API und MySQL, produktiv in EU West (Amsterdam).",
        "Google Cloud/Firebase: Cloud Messaging und technische Push-Metadaten; keine Datei- oder Datenbankspeicherung für SSD Manager.",
        "Apple und Google: Push-Infrastruktur und App-Store-Bereitstellung nach deren eigener Verantwortlichkeit beziehungsweise technischer Rolle.",
        "ALL-INKL.COM – Neue Medien Münnich: ausschließlich verschlüsselte Offsite-Datenbanksicherungen in einem getrennten FTP-Bereich.",
    ]), ("p", "Die jeweils aktuelle Unterauftragnehmerliste ist Bestandteil des AVV. Drittlandzugriffe werden nur unter den Voraussetzungen der Art. 44 ff. DSGVO zugelassen.")]),
    ("5. Speicherdauer und Löschung", [("bullets", [
        "Konten: zunächst Deaktivierung; nach autorisierter Löschvormerkung 30 Tage Karenz, danach Anonymisierung und Entfernung kontoidentifizierender Daten.",
        "Diensthistorie und Ankündigungstexte: bleiben für den schulischen Kontext erhalten, jedoch nach Kontolöschung ohne Namenszuordnung als „Gelöschter Nutzer“.",
        "Anhänge eines gelöschten Kontos: vollständige Löschung; im Chat bleibt nur ein Nichtverfügbarkeits-Hinweis.",
        "Nicht beanspruchte Uploads: automatische Löschung nach einem Tag.",
        "Abgeschlossene Inhaltsmeldungen: automatische Löschung nach 12 Monaten; offene Meldungen bleiben bis zur Bearbeitung verfügbar.",
        "Benachrichtigungs- und Loginprotokolle: 90 Tage, soweit kein Sicherheitsfall eine dokumentierte Sperre rechtfertigt.",
        "Verschlüsselte Backups: tägliche Rotation, Löschung nach 30 Tagen; keine Wiederaufnahme gelöschter Daten in den Regelbetrieb.",
        "Vertrags- und Rechnungsunterlagen: nach den gesetzlichen handels- und steuerrechtlichen Fristen.",
    ])]),
    ("6. Betroffenenrechte", [("p", "Betroffene haben nach Maßgabe der gesetzlichen Voraussetzungen Rechte auf Auskunft, Berichtigung, Löschung, Einschränkung, Datenübertragbarkeit und Widerspruch sowie auf Beschwerde bei einer Datenschutzaufsichtsbehörde. Anfragen zu schulischen Daten sind zunächst an Schule/Schulträger zu richten; MinutMate unterstützt die verantwortliche Stelle. Anfragen zu eigenen MinutMate-Verarbeitungen: datenschutz@minutmate.com.")]),
    ("7. Minderjährige, Bereitstellung und Automatisierung", [("p", "Konten werden ausschließlich durch autorisierte schulische Stellen eingerichtet; es gibt keine offene Registrierung. Ohne erforderliche Konto-, Rollen- und Authentifizierungsdaten ist die App nicht nutzbar. SSD Manager trifft keine ausschließlich automatisierten Entscheidungen mit rechtlicher oder ähnlich erheblicher Wirkung und verwendet keine Werbe- oder Tracking-SDKs.")]),
    ("8. Sicherheit und Aktualisierung", [("p", "Die wesentlichen Sicherheitsmaßnahmen sind in den TOMs beschrieben. Diese Erklärung ist bei Funktions-, Anbieter-, Rechts- oder Speicheränderungen zu aktualisieren. Vor Veröffentlichung sind Versionsdatum, Store-URLs und die verantwortliche Schule beziehungsweise ein schulbezogener Hinweis einzusetzen.")]),
])
register("02_SSD_Manager_Datenschutzerklaerung.docx", "Datenschutzerklärung", "Öffentliche App-Datenschutzinformation", "Schulrechtsgrundlage/URLs prüfen")

save("03_SSD_Manager_Nutzungsbedingungen.docx", "Nutzungsbedingungen SSD Manager", "Regeln für autorisierte schulische Nutzende", [
    ("1. Geltungsbereich", [("p", f"Anbieter ist {OPERATOR}, {ADDRESS}. Die App wird ausschließlich in von MinutMate freigeschalteten Schulumgebungen genutzt. Für Schulen gelten ergänzend SaaS-Vertrag, AVV, TOMs und individuelle Leistungsbeschreibungen.")]),
    ("2. Leistung", [("p", "SSD Manager unterstützt Dienstplanung, Diensthistorie, Ankündigungen, Anhänge, Rollen, Konten, Geräte und schulische Administration des Schulsanitätsdienstes. Die App ersetzt weder medizinische Dokumentation noch Notruf-, Einsatzleit- oder Patientenakten-Systeme.")]),
    ("3. Konten und Minderjährige", [("bullets", ["Keine Selbstregistrierung; Konten werden nach schulischer Freigabe angelegt.", "Zugangsdaten sind persönlich und vertraulich zu behandeln.", "Schule/Schulträger entscheidet über Rechtsgrundlage, Nutzerkreis, Information und erforderliche Beteiligungen.", "Unangemessene Ankündigungen können in der App an Sani-Leitung und Lehreraufsicht gemeldet werden; weitere Verdachtsfälle sind unverzüglich an Schule oder support@ssd-manager.minutmate.com zu melden."])]),
    ("4. Zulässige Nutzung", [("bullets", ["Nur schulorganisatorische Zwecke des Schulsanitätsdienstes.", "Keine Patientendaten, Diagnosen, Behandlungsprotokolle oder sonstige medizinische Falldokumentation im Ankündigungschat.", "Keine rechtswidrigen, beleidigenden, diskriminierenden oder fremde Rechte verletzenden Inhalte.", "Keine Umgehung von Rollen, Sicherheitsmaßnahmen, Speichergrenzen oder Mandantentrennung."])]),
    ("5. Inhalte und Administration", [("p", "Nutzende behalten Rechte an rechtmäßig bereitgestellten Inhalten und räumen die zur Speicherung, Anzeige, Übermittlung und Sicherung erforderlichen Nutzungsrechte für die Vertragsdauer ein. Autorisierte Lehreraufsicht und Sani-Leitung dürfen gemeldete Inhalte prüfen, als unbedenklich schließen, entfernen und bei Missbrauch den verantwortlichen Account sperren.")]),
    ("6. Verfügbarkeit und Support", [("p", "MinutMate bemüht sich um zuverlässigen Betrieb; Wartung, Sicherheitsmaßnahmen und Anbieter- oder Netzstörungen können Unterbrechungen verursachen. Support: support@ssd-manager.minutmate.com. Garantierte Service-Level gelten nur, wenn sie ausdrücklich vertraglich vereinbart wurden.")]),
    ("7. Datenschutz, Sperrung und Beendigung", [("p", "Datenschutzinformationen, AVV und TOMs gelten ergänzend. Konten können bei Sicherheitsrisiken, Schulweisung oder Vertragsverstoß gesperrt werden. Bei Vertragsende gelten Export-, Anonymisierungs- und Löschregeln des SaaS-Vertrags und AVV.")]),
    ("8. Haftung und Schlussbestimmungen", [("p", "Es gelten die gesetzlichen Haftungsregeln. Für Vorsatz, grobe Fahrlässigkeit, Personenschäden und zwingende gesetzliche Haftung wird unbeschränkt gehaftet. Bei leicht fahrlässiger Verletzung wesentlicher Vertragspflichten ist die Haftung auf den vorhersehbaren, vertragstypischen Schaden begrenzt. Es gilt deutsches Recht; zwingende Schutzvorschriften bleiben unberührt.")]),
])
register("03_SSD_Manager_Nutzungsbedingungen.docx", "Nutzungsbedingungen", "Regeln innerhalb der App", "Haftung anwaltlich prüfen")

save("04_SSD_Manager_SaaS_Vertrag_Schule.docx", "SaaS-Nutzungsvertrag SSD Manager", "Vertrag über die Bereitstellung für Schulen", [
    ("Vertragspartner", [("table", ["Partei", "Angaben"], [("Schule / Schulträger", "Name, Anschrift, Vertretung und E-Mail: ______________________________________________"), ("Anbieter", f"{OPERATOR}\n{ADDRESS}\nverwaltung@minutmate.com")], [2200, 7160])]),
    ("1. Vertragsgegenstand", [("p", "MinutMate stellt einen mandantengetrennten SSD-Manager-Workspace als Software-as-a-Service bereit. Der vereinbarte Umfang umfasst Dienstplanung, Rollen/Konten, Ankündigungen, Anhänge, Geräteverwaltung und Administration.")]),
    ("2. Vertragsunterlagen und Rangfolge", [("numbered", ["Individuelles Angebot und dieser Vertrag.", "AVV einschließlich Unterauftragnehmerliste und TOMs.", "Leistungs-/Supportvereinbarung, soweit vereinbart.", "Nutzungsbedingungen."])]),
    ("3. Einrichtung und Mitwirkung", [("bullets", ["Die Schule benennt verantwortliche und technische Kontakte.", "Die Schule bestimmt Rechtsgrundlage, Nutzerkreis, Rollen und zulässige Inhalte.", "Kontodaten werden ausschließlich über autorisierte schulische Anfragen bereitgestellt.", "Administrationskonten, Initialpasswörter und Endgeräte werden angemessen geschützt.", "Sicherheits- und Datenschutzvorfälle werden unverzüglich gemeldet."])]),
    ("4. Leistung und Support", [("p", "Produktsupport erfolgt über support@ssd-manager.minutmate.com. Ziel ist eine erste Reaktion innerhalb von ein bis drei Werktagen. Ein verbindliches SLA besteht nur bei ausdrücklicher Vereinbarung.")]),
    ("5. Vergütung", [("table", ["Feld", "Vereinbarung"], [("Lizenzpreis", "________________ EUR"), ("Abrechnungszeitraum", "________________"), ("Zahlungsziel", "14 Tage ab Rechnungseingang"), ("Vertragsbeginn", "________________")], [2600, 6760])]),
    ("6. Datenschutz und Sicherheit", [("p", "Vor Verarbeitung echter Schuldaten schließen die Parteien den AVV. Die Schule bleibt für Zweck, Rechtsgrundlage, Transparenz, Datenminimierung, Rollen und schulrechtliche Zulässigkeit verantwortlich. MinutMate erfüllt die vereinbarten TOMs.")]),
    ("7. Laufzeit und Kündigung", [("p", "Laufzeit: [ ] Schulsemester  [ ] Schuljahr  [ ] abweichend: __________________. Ordentliche Kündigungsfrist: 14 Tage zum Laufzeitende, soweit individuell nichts anderes vereinbart ist. Außerordentliche Kündigung aus wichtigem Grund bleibt möglich.")]),
    ("8. Vertragsende, Export und Löschung", [("p", "Nach Vertragsende wird die Schulumgebung deaktiviert. Die Schule kann innerhalb von 30 Tagen einen technisch angemessenen Export anfordern. Produktive Schuldaten werden spätestens 30 Tage nach Vertragsende nach dokumentierter Abstimmung gelöscht oder anonymisiert; Backups rotieren innerhalb weiterer 30 Tage aus. Gesetzlich aufzubewahrende Vertrags- und Abrechnungsdaten bleiben getrennt erhalten.")]),
    ("9. Mängel, Haftung und Vertraulichkeit", [("p", "Reproduzierbare Mängel sind nachvollziehbar zu melden. MinutMate behebt erhebliche Mängel innerhalb angemessener Frist. Haftung richtet sich nach den gesetzlichen Regeln und der individuell geprüften Vertragsfassung. Nicht öffentliche technische, geschäftliche, schulische und personenbezogene Informationen werden vertraulich behandelt.")]),
    ("10. Schlussbestimmungen", [("p", "Änderungen bedürfen mindestens der Textform. Es gilt deutsches Recht. Unwirksame Einzelbestimmungen lassen den Vertrag im Übrigen unberührt.")]),
], sign=True)
register("04_SSD_Manager_SaaS_Vertrag_Schule.docx", "SaaS-Vertrag", "Hauptvertrag mit Schule/Schulträger", "Preis/Laufzeit und Recht prüfen")

save("05_SSD_Manager_AVV_Schule.docx", "Auftragsverarbeitungsvertrag SSD Manager", "Vereinbarung nach Art. 28 DSGVO", [
    ("Vertragspartner und Rollen", [("table", ["Partei", "Rolle und Angaben"], [("Schule / Schulträger", "Verantwortlicher; Name/Anschrift/Vertretung/Kontakt: ______________________________"), ("MinutMate", f"Auftragsverarbeiter; {OPERATOR}, {ADDRESS}, datenschutz@minutmate.com")], [2200, 7160]), ("p", "Eigene Verarbeitungen von MinutMate für Vertragsverwaltung, Abrechnung, IT-Sicherheit und gesetzliche Pflichten erfolgen in eigener Verantwortlichkeit.")]),
    ("1. Gegenstand, Art, Zweck und Dauer", [("p", "Bereitstellung, Hosting und Administration einer schulbezogenen SSD-Manager-Umgebung; Konten, Rollen, Dienstplanung, Ankündigungen, Anhänge, Geräte/Sitzungen, Support, Sicherung und IT-Sicherheit. Dauer des Hauptvertrags einschließlich abgestimmter Export- und Löschphase.")]),
    ("2. Betroffene Personen und Daten", [("table", ["Kategorie", "Inhalt"], [("Betroffene", "Schülerinnen und Schüler im Schulsanitätsdienst, Sani-Leitung, Lehreraufsicht, Sekretariat, sonstige autorisierte schulische Ansprechpartner"), ("Stammdaten", "Name, Benutzername, E-Mail, Schule, Rolle, Status, Qualifikationsbeginn"), ("Organisationsdaten", "Diensttage, Zuweisungen, Krankmeldungen, Historie"), ("Kommunikation", "Ankündigungstexte, Absender, Fotos, Dateien und Metadaten"), ("Technik/Sicherheit", "Sitzungs- und Geräteinformationen, Installationskennung, Push-Token, Audit-, Login- und Fehlerdaten"), ("Support", "Anfragen und zur Problemlösung übermittelte Informationen")], [2400, 6960]), ("p", "Medizinische Falldaten und besondere Kategorien personenbezogener Daten sind nicht planmäßig vorgesehen. Sie dürfen nicht in Ankündigungen oder Anhängen gespeichert werden, sofern keine ausdrückliche Weisung und gesonderte rechtliche Prüfung vorliegt.")]),
    ("3. Weisungsbindung", [("p", "MinutMate verarbeitet Schuldaten ausschließlich auf dokumentierte Weisung. Hauptvertrag, AVV, berechtigte Konfigurationen und nachweisbare Mitteilungen gelten als Weisungen. Offensichtlich rechtswidrige Weisungen werden beanstandet und bis zur Klärung ausgesetzt, soweit rechtlich zulässig.")]),
    ("4. Pflichten des Auftragsverarbeiters", [("bullets", ["Vertraulich verpflichtete, erforderliche Zugriffsberechtigte.", "Umsetzung und fortlaufende Prüfung der vereinbarten TOMs.", "Unterstützung bei Betroffenenrechten, DSFA, Konsultationen, Sicherheitsprüfungen und Meldungen.", "Unverzügliche Information der Schule nach Bekanntwerden einer Datenschutzverletzung.", "Führen erforderlicher Nachweise und Bereitstellung angemessener Kontrollinformationen."])]),
    ("5. Unterauftragsverarbeiter und Drittlandtransfers", [("p", "Die separate Unterauftragnehmerliste ist allgemein genehmigt. MinutMate informiert über beabsichtigte wesentliche Änderungen mit angemessener Frist; die Schule kann aus wichtigem datenschutzrechtlichem Grund widersprechen. Unterauftragsverarbeiter werden gleichwertig verpflichtet. Übermittlungen außerhalb des EWR erfolgen nur nach Art. 44 ff. DSGVO.")]),
    ("6. Kontrollen", [("p", "MinutMate stellt erforderliche Informationen und Nachweise bereit. Dokumentenprüfungen haben Vorrang. Erforderliche Vor-Ort- oder Fernkontrollen erfolgen abgestimmt, während üblicher Geschäftszeiten, mit angemessener Vorankündigung und unter Wahrung von Sicherheits- und Geheimhaltungsinteressen.")]),
    ("7. Rückgabe und Löschung", [("p", "Nach Vertragsende wird der Schulzugang deaktiviert. Exportfrist: 30 Tage. Danach werden produktive Schuldaten nach Weisung gelöscht oder anonymisiert. Verschlüsselte tägliche Sicherungen rotieren spätestens 30 Tage nach Erstellung aus und bleiben bis dahin gesperrt; sie werden nicht zur Rücknahme einer wirksamen Löschung in den Regelbetrieb verwendet. Gesetzliche Aufbewahrungspflichten für eigene Vertragsdaten bleiben unberührt.")]),
    ("8. Anlagen", [("bullets", ["Anlage 1: Verarbeitungsbeschreibung und Datenkategorien (Abschnitte 1 und 2).", "Anlage 2: Unterauftragnehmerliste SSD Manager.", "Anlage 3: Technische und organisatorische Maßnahmen.", "Anlage 4: Lösch- und Aufbewahrungskonzept."])]),
], sign=True)
register("05_SSD_Manager_AVV_Schule.docx", "AVV", "Pflichtvertrag bei Auftragsverarbeitung", "Art.-28-Prüfung und Unterschrift")

save("06_SSD_Manager_TOMs.docx", "Technische und organisatorische Maßnahmen", "SSD Manager – Art. 28 Abs. 3 und Art. 32 DSGVO", [
    ("1. Governance und Risikoorientierung", [("bullets", ["Mindestens halbjährliche sowie anlassbezogene Überprüfung.", "Berechtigungs- und Restore-Prüfungen mindestens halbjährlich; erster Restore-Test erfolgreich durchgeführt.", "Geheimnisse ausschließlich in geschützten Produktionsvariablen und lokalen, Git-ausgeschlossenen Recovery-Ablagen.", "Vertraulichkeitsverpflichtung für Personen mit administrativem Zugriff."])]),
    ("2. Physische und lokale Sicherheit", [("bullets", ["Produktive Systeme in vertraglich gebundenen Rechenzentren der Dienstleister.", "Lokale Administrationsgeräte mit Betriebssystemanmeldung, Gerätesperre, Festplattenverschlüsselung, Updates und Malware-Schutz.", "Keine produktiven Daten auf privaten Wechselmedien; temporäre Restore-Dateien werden nach Prüfung gelöscht."])]),
    ("3. Authentifizierung und Zugang", [("bullets", ["Passwort-Hashing; keine Klartextpasswörter in der Datenbank.", "Kurzlebige Zugriffstoken und widerrufbare, serverseitig gehashte Refresh-Sitzungen.", "Persistente zufällige Installationskennung ersetzt veraltete Sitzungen derselben Installation.", "Login-Rate-Limits und sichere, generische Fehlermeldungen.", "Deaktivierung und Löschvormerkung widerrufen alle Sitzungen."])]),
    ("4. Rollen, Zugriff und Mandantentrennung", [("bullets", ["Serverseitige Schulgrenzen über school_id und rollenbezogene Autorisierung.", "Rollen Schulsanitäter, Sani-Leitung, Lehreraufsicht und Sekretariat mit minimalen Rechten.", "Kontolöschung nur durch Lehreraufsicht oder Sani-Leitung.", "Dateizugriffe authentifiziert, uploader- und schulbezogen; 100-MB-Nutzerquote und maximal vier Anhänge à 8 MB je Nachricht.", "Berechtigungen werden bei Rollenwechsel, Ausscheiden und Vertragsende entzogen."])]),
    ("5. Verschlüsselung und Übertragung", [("bullets", ["Produktive API ausschließlich über HTTPS/TLS.", "Datenbank nicht öffentlich exponiert; API und MySQL im selben privaten Railway-Projekt und in derselben EU-Region.", "Tägliche SQL-Sicherungen werden komprimiert und mit XChaCha20-Poly1305 authentifiziert verschlüsselt; All-Inkl erhält keinen Schlüssel.", "FTPS-Übertragung mit Zertifikatsprüfung; dedizierter, auf /ssd-manager-backups/ beschränkter FTP-Nutzer."])]),
    ("6. Protokollierung und Datenminimierung", [("bullets", ["Auditierung administrativer und sicherheitsrelevanter Vorgänge.", "Keine Roh-Exceptions, SQL-, Token-, Passwort- oder Antwortinhalte in Nutzerfehlermeldungen.", "FCM-Payloads enthalten nur erforderliche technische Routinginformationen.", "Keine Werbung, Marketinganalyse oder offene Registrierung.", "Medizinische Falldokumentation ist organisatorisch und vertraglich ausgeschlossen."])]),
    ("7. Verfügbarkeit, Backup und Wiederherstellung", [("bullets", ["Railway API und MySQL in EU West (Amsterdam).", "Separater 15-Minuten-Wartungsjob für Statusübergänge, Erinnerungen, verwaiste Uploads, Löschungen und Logrotation.", "Tägliches Offsite-Backup um 03:35 Uhr Europe/Berlin; 30-Tage-Rotation.", "Erster realer Backup-, Download-, Authentifizierungs-, Entschlüsselungs- und gzip-Test erfolgreich.", "Zielwerte RPO/RTO jeweils 24 Stunden; Abweichungen und Wiederanlauf werden dokumentiert."])]),
    ("8. Incident Response und Überprüfung", [("numbered", ["Vorfall erkennen, eindämmen und Beweismittel zugriffsgeschützt sichern.", "Betroffene Systeme, Daten, Schulen, Personen und Zeiträume bestimmen.", "Schlüssel/Tokens rotieren und unberechtigte Sitzungen widerrufen.", "Betroffene verantwortliche Schule unverzüglich mit verfügbaren Informationen informieren.", "Ursache beheben, kontrolliert wiederherstellen und Wirksamkeit prüfen.", "Entscheidungen, Meldungen und Verbesserungen dokumentieren."])]),
])
register("06_SSD_Manager_TOMs.docx", "TOMs", "Sicherheitsanlage zum AVV", "Anbieter-TOMs beilegen/regelmäßig prüfen")

save("07_SSD_Manager_Unterauftragnehmerliste.docx", "Unterauftragnehmerliste SSD Manager", "Aktive Dienstleister der Auftragsverarbeitung", [
    ("Geltungsbereich", [("p", "Diese Liste ist Anlage des SSD-Manager-AVV. Sie enthält nur im SSD-Manager-Produktivbetrieb aktive oder für dessen Sicherung eingesetzte Dienstleister. StudyConnect-spezifische Anbieter werden nicht übernommen.")]),
    ("Aktive Unterauftragnehmer", [("table", ["Anbieter", "Zweck / Daten", "Region / Transfer", "Nachweisstatus"], [
        ("Railway Corporation", "API, MySQL, Logs; Konto-, Organisations-, Kommunikations-, Datei- und Sicherheitsdaten", "EU West, Amsterdam; Anbieterzugriffe/Unterauftragnehmer nach DPA prüfen", "Railway DPA und signierte Kopie vorhanden; Kontozuordnung jährlich belegen"),
        ("Google Cloud / Firebase", "FCM/APNs-Push; Geräte-/Push-Token und technische Nachrichtenmetadaten", "Internationale Infrastruktur möglich; DPA/Transfermechanismus prüfen", "Google/Firebase-Vertragsnachweis vorhanden; SSD-Projektzuordnung belegen"),
        ("ALL-INKL.COM – Neue Medien Münnich", "Ausschließlich verschlüsselte Offsite-Datenbanksicherungen", "Deutschland; Inhalte ohne separaten Schlüssel nicht lesbar", "AVV aus MinutMate-Vertragsakte vorhanden; Leistungsbezug SSD Manager dokumentieren"),
    ], [1800, 3100, 2400, 2060])]),
    ("Nicht als SSD-Unterauftragnehmer übernommen", [("bullets", ["Vercel: kein aktiver SSD-Manager-Produktionsdienst.", "Resend: kein aktiver SSD-Manager-Mailversanddienst.", "Firebase Storage/Google Cloud Storage: keine SSD-Dateiablage.", "GitHub: Entwicklungs- und Deploymentquelle, nicht regulärer Speicher produktiver Schuldaten.", "Apple/Google App Stores: gesondert nach konkreter Rolle und eigener Verantwortlichkeit dokumentieren."])]),
    ("Änderungsverfahren", [("p", "MinutMate informiert Vertragsschulen vor einer wesentlichen Änderung entsprechend dem AVV. Ein Widerspruch ist aus wichtigem datenschutzrechtlichem Grund möglich. DPA, Region, Transfergrundlage, Unterauftragnehmer und Kontozuordnung werden mindestens jährlich und anlassbezogen geprüft.")]),
])
register("07_SSD_Manager_Unterauftragnehmerliste.docx", "Unterauftragnehmerliste", "AVV-Anlage und Transparenz", "DPA-Zuordnungen final belegen")

save("08_SSD_Manager_Loesch_und_Aufbewahrungskonzept.docx", "Lösch- und Aufbewahrungskonzept", "Verbindliche Regeln für Live-System, Protokolle und Sicherungen", [
    ("Grundsätze", [("p", "Löschung wird durch Schule/Schulträger als Verantwortlichen veranlasst. Lehreraufsicht oder Sani-Leitung darf nach schulischer Prüfung eine Kontolöschung vormerken. Die 30-tägige Karenz ermöglicht Korrektur; anschließend erfolgt die automatisierte Verarbeitung.")]),
    ("Fristenmatrix", [("table", ["Datenart", "Auslöser", "Frist / Ergebnis"], [
        ("Aktives Konto", "Ausscheiden/Schulweisung", "Sofort deaktivieren; alle Sitzungen widerrufen"),
        ("Vorgemerkte Kontolöschung", "Autorisierte Vormerkung", "30 Tage; dann Identifikatoren entfernen und Konto anonymisieren"),
        ("Diensthistorie", "Kontolöschung", "Historisch erhalten; Namensbezug durch „Gelöschter Nutzer“ ersetzen"),
        ("Ankündigungstext", "Kontolöschung", "Inhalt erhalten; Absender anonymisieren"),
        ("Gemeldeter Inhalt", "Moderationsentscheidung", "Text und Dateibytes entfernen; Tombstones im Chat erhalten"),
        ("Inhaltsmeldung", "Abschluss", "Nach 12 Monaten automatisch löschen"),
        ("Anhänge", "Uploader-Löschung/Schulende", "Dateibytes vollständig löschen; Tombstone im Chat"),
        ("Unbeanspruchte Uploads", "Upload ohne Nachricht", "Nach 1 Tag automatisch löschen"),
        ("Sitzungen/Push-Token", "Logout/Widerruf/Deaktivierung", "Sofort widerrufen oder entfernen"),
        ("Login-/Benachrichtigungslogs", "Ereignis", "90 Tage, außer dokumentierter Sicherheits-/Rechtsfall"),
        ("Auditnachweise", "Ereignis", "Regelfrist 12 Monate; längere Sperre nur dokumentiert"),
        ("Backups", "Erstellung", "Täglich; automatische Remote-Löschung nach 30 Tagen"),
        ("Schulumgebung", "Vertragsende", "Deaktivierung; 30 Tage Export; danach produktive Löschung/Anonymisierung"),
        ("Vertrags/Rechnungsdaten", "Vertragsende", "Gesetzliche handels- und steuerrechtliche Fristen"),
        ("Betroffenenanfragen", "Abschluss", "Nachweis grundsätzlich 3 Jahre, datensparsam"),
    ], [2200, 2300, 4860])]),
    ("Backups und Wiederherstellung", [("p", "Wirksame Löschungen werden nicht allein zur Wiederherstellung einzelner Konten aus Backups rückgängig gemacht. Bei einer vollständigen Notfallwiederherstellung sind seit Sicherungszeitpunkt wirksam gewordene Löschungen anhand geschützter Löschprotokolle erneut anzuwenden. Backups sind verschlüsselt, zweckgebunden, gesperrt und nach 30 Tagen technisch entfernt.")]),
    ("Nachweis und Ausnahmen", [("bullets", ["Automatisierte Jobs und Stichproben werden dokumentiert.", "Löschfehler werden als Vorfall bearbeitet.", "Rechtliche Aufbewahrungssperren enthalten Grund, Umfang, Genehmigung und Enddatum.", "Fristen werden mindestens jährlich und bei Produktänderungen geprüft."])]),
])
register("08_SSD_Manager_Loesch_und_Aufbewahrungskonzept.docx", "Löschkonzept", "Interner Nachweis und AVV-Anlage", "Schulfreigabe erforderlich")

save("09_SSD_Manager_DSFA_Vorpruefung_und_Muster.docx", "DSFA-Vorprüfung und Muster", "Unterstützung der verantwortlichen Schule nach Art. 35 DSGVO", [
    ("Verantwortung und Ergebnis der Vorprüfung", [("p", "Die endgültige DSFA-Entscheidung und Durchführung liegen bei Schule/Schulträger. Wegen systematischer schulischer Organisation, minderjähriger Betroffener, Rollen-/Diensthistorie und Kommunikationsinhalten ist mindestens eine dokumentierte Schwellwertanalyse erforderlich. Eine vollständige DSFA ist je nach Landesrecht, Umfang und konkreter Nutzung ernsthaft in Betracht zu ziehen.")]),
    ("Beschreibung der Verarbeitung", [("p", "Zweck: Organisation des Schulsanitätsdienstes. Vorgänge: Anlage schulischer Konten, Rollen, Dienstplanung, Krankmeldungen, Ankündigungen/Anhänge, Geräte/Sitzungen, Push, Support, Sicherung und Löschung. Keine offene Registrierung, keine Werbung, keine Leistungsbewertung, keine medizinische Falldokumentation.")]),
    ("Notwendigkeit und Verhältnismäßigkeit", [("bullets", ["Nur erforderliche Rollen und Datenfelder; Schule steuert Nutzerkreis.", "Schul- und Rollenprüfung serverseitig.", "Ankündigungskanal statt privater Chats.", "Keine planmäßige Verarbeitung besonderer Kategorien; verbindliches Inhaltsverbot für Falldaten.", "30-Tage-Karenz, Anonymisierung historischer Einträge und vollständige Dateilöschung.", "Betroffenenprozesse laufen über die verantwortliche Schule."])]),
    ("Risikobewertung", [("table", ["Risiko", "Auswirkung", "Maßnahmen", "Restrisiko"], [
        ("Unberechtigter schulübergreifender Zugriff", "Offenlegung Minderjährigendaten", "school_id-Prüfung, Rollen, API-Autorisierung, Tests", "niedrig/mittel"),
        ("Kontenmissbrauch", "Änderung von Diensten/Kommunikation", "Passwort-Hashing, Sitzungswiderruf, Rate-Limit, Geräteübersicht", "mittel"),
        ("Unzulässige medizinische Inhalte", "Art.-9-Daten im Chat/Dateien", "Inhaltsverbot, Schulregeln, Rollen, Löschung und Schulung", "mittel"),
        ("Verlust/Offenlegung von Anhängen", "Privatsphäre- und Reputationsschaden", "Authentifizierter Abruf, Quote, Größenlimit, Löschung", "mittel"),
        ("Push-Metadaten/Drittlandzugriff", "Technische Zuordnung", "Minimierte Payloads, DPA/Transferprüfung, keine Falldaten", "niedrig/mittel"),
        ("Datenverlust", "Ausfall der Dienstplanung/Historie", "EU-Hosting, verschlüsselte tägliche Backups, Restore-Test", "niedrig/mittel"),
        ("Überlange Speicherung", "Unverhältnismäßige Historie", "Fristen, 15-Minuten-Job, 30-Tage-Backuprotation, Jahresreview", "niedrig/mittel"),
    ], [2100, 2200, 3400, 1660])]),
    ("Von der Schule auszufüllen", [("bullets", ["Konkrete Rechtsgrundlage und einschlägiges Landes-Schulrecht.", "Zwecke, Nutzerzahlen, Altersgruppen, Standorte und Laufzeit.", "Stellungnahme des schulischen/behördlichen Datenschutzbeauftragten.", "Bewertung landesspezifischer Positiv-/Negativlisten.", "Verbleibende Risiken, Freigabeentscheidung und Überprüfungsdatum.", "Gegebenenfalls vorherige Konsultation der Aufsichtsbehörde nach Art. 36 DSGVO."])]),
])
register("09_SSD_Manager_DSFA_Vorpruefung_und_Muster.docx", "DSFA-Muster", "Risikoprüfung der Schule unterstützen", "Je Schule vervollständigen")

save("10_SSD_Manager_Verzeichnis_Verarbeitungstaetigkeiten.docx", "Verzeichnis von Verarbeitungstätigkeiten", "Muster nach Art. 30 DSGVO", [
    ("Hinweise", [("p", "Dieses Muster trennt eigene MinutMate-Verarbeitungen von der Verarbeitung im Auftrag. Schule/Schulträger führt ein eigenes Verzeichnis für die schulische Nutzung. Die Einträge sind auf tatsächliche Rechtsgrundlagen, Empfänger und Fristen anzupassen.")]),
    ("Verarbeitungstätigkeiten", [("table", ["Tätigkeit", "Rolle/Zweck", "Daten/Betroffene", "Empfänger/Frist"], [
        ("Schul-Workspace", "Auftragsverarbeitung; Dienstplanung und Kommunikation", "Konten, Rollen, Dienste, Nachrichten, Dateien; schulische Nutzende", "Railway/Firebase/All-Inkl; nach Schulweisung und Löschkonzept"),
        ("Vertragsverwaltung", "Eigene Verantwortung; Vertrag, Abrechnung, Nachweis", "Schulkontakte, Vertrags- und Rechnungsdaten", "Steuer-/Beratungsstellen nach Bedarf; gesetzliche Fristen"),
        ("IT-Sicherheit", "Eigene/auftragsbezogene Sicherheit; Art. 6 Abs. 1 f bzw. Art. 32", "IP-/Login-/Audit-/Geräte- und Fehlerdaten", "Railway; Regel 90 Tage, Audit 12 Monate"),
        ("Support", "Vertragserfüllung und Weisungsunterstützung", "Kontakt, Anfrage, erforderliche technische Daten", "Autorisierte MinutMate-Personen; Abschluss plus angemessene Nachweisfrist"),
        ("Datenschutzanfragen", "Art. 6 Abs. 1 c; Rechtebearbeitung/Nachweis", "Identität, Anfrage, Kommunikation, Entscheidung", "Schule/Betroffene/Aufsicht; grundsätzlich 3 Jahre"),
        ("App-Store-Bereitstellung", "Distribution, Sicherheit und Updates", "Entwicklerkonto-, Store- und technische Abrufdaten", "Apple/Google nach deren Bedingungen; anbieterabhängig"),
    ], [1800, 2400, 2800, 2360])]),
    ("Pflichtfelder zur Finalisierung", [("bullets", ["Konkrete Rechtsgrundlagen und Interessenabwägungen.", "Tatsächliche Empfänger, Drittlandgarantien und DPA-Versionen.", "Kategorien von Löschsperren und Verantwortliche.", "Name/Kontakt des Datenschutzbeauftragten, falls bestellt.", "Datum der Freigabe und nächste Überprüfung."])]),
])
register("10_SSD_Manager_Verzeichnis_Verarbeitungstaetigkeiten.docx", "VVT", "Interne Rechenschaftspflicht", "MinutMate und Schule getrennt finalisieren")

save("11_SSD_Manager_Betroffenenanfragen_Prozess.docx", "Prozess für Betroffenenanfragen", "Auskunft, Berichtigung, Löschung und weitere Rechte", [
    ("Grundsatz", [("p", "Schulische Anfragen werden durch Schule/Schulträger entschieden. MinutMate trifft als Auftragsverarbeiter keine eigenständige materielle Entscheidung, unterstützt jedoch technisch. Direkte Anfragen an MinutMate werden unverzüglich zuständigkeitshalber weitergeleitet; eigene MinutMate-Verarbeitungen werden direkt bearbeitet.")]),
    ("Ablauf", [("numbered", ["Eingang mit Datum, Kanal und Begehren erfassen; keine unnötigen Ausweiskopien anfordern.", "Zuständigkeit und Identität risikoorientiert prüfen.", "Frist notieren; grundsätzlich unverzüglich und spätestens innerhalb eines Monats reagieren.", "Datenquellen, Empfänger, Fristen und Rechteausnahmen bestimmen.", "Export, Berichtigung, Einschränkung, Anonymisierung oder Löschung technisch ausführen beziehungsweise der Schule bereitstellen.", "Antwort sicher übermitteln; Ablehnungen und Fristverlängerungen begründen.", "Abschluss und Nachweis datensparsam dokumentieren; regulär drei Jahre."])]),
    ("Technische SSD-Manager-Unterstützung", [("bullets", ["Manager können ein Nutzerdaten-Exportarchiv erzeugen.", "Deaktivierung sperrt Zugang und widerruft Sitzungen sofort.", "Löschvormerkung startet eine 30-Tage-Karenz.", "Historische Dienste/Nachrichten werden anonymisiert; Anhänge vollständig gelöscht.", "Backups rotieren innerhalb von 30 Tagen aus; Notfallrestores wenden Löschungen erneut an."])]),
    ("Kontakte und Protokollfelder", [("p", "Produkt: support@ssd-manager.minutmate.com · Datenschutz MinutMate: datenschutz@minutmate.com · Schule/DSB: __________________. Aktenzeichen, Eingang, Identitätsprüfung, Verantwortlicher, Maßnahmen, Freigabe, Antwortdatum und Löschdatum sind zu dokumentieren.")]),
])
register("11_SSD_Manager_Betroffenenanfragen_Prozess.docx", "Betroffenenprozess", "Operativer Rechteprozess", "Schulkontakte ergänzen")

save("12_SSD_Manager_Datenschutzverletzung_Prozess.docx", "Prozess bei Datenschutzverletzungen", "Erkennung, Eindämmung, Information und Nachbereitung", [
    ("Meldewege", [("p", "Interner/technischer Meldeweg: support@ssd-manager.minutmate.com und verwaltung@minutmate.com. Datenschutzkontakt: datenschutz@minutmate.com. Notfallkontakt der Schule/Schulträger: __________________. MinutMate informiert die verantwortliche Stelle unverzüglich nach Bekanntwerden eines Vorfalls im Auftragsbereich.")]),
    ("Sofortmaßnahmen", [("numbered", ["Vorfallzeitpunkt, Meldende, Systeme und erste Beobachtungen dokumentieren.", "Unbefugte Sitzungen sperren; Tokens, Schlüssel oder Passwörter risikoorientiert rotieren.", "Betroffene Daten, Schulen, Personen, Dauer und mögliche Exfiltration bestimmen.", "Beweismittel zugriffsgeschützt sichern; keine personenbezogenen Details in allgemeine Chats/Logs kopieren.", "Verantwortliche Schule unverzüglich mit verfügbaren Angaben informieren.", "Technische Ursache beheben und Systeme kontrolliert wieder in Betrieb nehmen."])]),
    ("Bewertung und Meldung", [("p", "Schule/Schulträger bewertet Risiko und entscheidet über Meldung an die Aufsichtsbehörde binnen 72 Stunden sowie gegebenenfalls Information Betroffener. MinutMate liefert fortlaufend Art des Vorfalls, Kategorien und ungefähre Zahl Betroffener/Datensätze, wahrscheinliche Folgen, Maßnahmen und Kontaktstelle. Eigene gesetzliche Pflichten von MinutMate bleiben unberührt.")]),
    ("Nachbereitung", [("bullets", ["Ursachenanalyse und Wirksamkeitskontrolle.", "Dokumentation auch bei Entscheidung gegen eine Meldung.", "TOMs, Tests, Schulung, Verträge und DSFA anpassen.", "Fristgebundene Löschung der Vorfallakte nach Rechts-/Nachweisprüfung."])]),
])
register("12_SSD_Manager_Datenschutzverletzung_Prozess.docx", "Incident-Prozess", "72-Stunden- und Auftragsmeldeprozess", "Notfallkontakte ergänzen")

save("13_SSD_Manager_Vertraulichkeitsverpflichtung.docx", "Vertraulichkeitsverpflichtung SSD Manager", "Für Personen mit administrativem oder Supportzugriff", [
    ("Verpflichtete Person", [("table", ["Feld", "Angabe"], [("Name", "____________________________________________"), ("Funktion/Tätigkeit", "____________________________________________"), ("Beginn", "____________________________________________")], [2500, 6860])]),
    ("Pflichten", [("bullets", ["Personenbezogene, schulische, technische und geschäftliche Informationen nur im erforderlichen Aufgaben- und Weisungsumfang verarbeiten.", "Keine unbefugte Offenlegung, private Nutzung, Kopie oder Weitergabe.", "Passwörter, Schlüssel, Tokens und Geräte wirksam schützen.", "Nur genehmigte Systeme und Übertragungswege verwenden.", "Fehlversand, Verlust, unberechtigten Zugriff und Schadsoftware unverzüglich melden.", "Unterlagen und lokale Daten bei Tätigkeitsende zurückgeben oder sicher löschen."])]),
    ("Fortgeltung und Bestätigung", [("p", "Die Verpflichtung gilt auch nach Ende der Tätigkeit. Verstöße können vertragliche, arbeitsrechtliche, schadensersatzrechtliche, ordnungswidrigkeiten- oder strafrechtliche Folgen haben. Die Person bestätigt, diese Pflichten und die SSD-Manager-Sicherheitsvorgaben verstanden zu haben.")]),
], sign=True)
register("13_SSD_Manager_Vertraulichkeitsverpflichtung.docx", "Vertraulichkeitsverpflichtung", "Personal-/Supportnachweis", "Bei Zugriff unterschreiben")

save("14_SSD_Manager_Schul_Onboarding_Freigabe.docx", "Schul-Onboarding und Freigabe", "Checkliste vor Aktivierung einer Schulumgebung", [
    ("Kontakte und Zuständigkeit", [("table", ["Rolle", "Name / Kontakt / Freigabe"], [("Schule/Schulträger Verantwortlicher", "____________________________________________"), ("Schulischer Datenschutzbeauftragter", "____________________________________________"), ("Lehreraufsicht", "____________________________________________"), ("Sani-Leitung", "____________________________________________"), ("Technischer Kontakt", "____________________________________________"), ("MinutMate", "info@ssd-manager.minutmate.com / support@ssd-manager.minutmate.com")], [3300, 6060])]),
    ("Pflichtprüfungen", [("bullets", ["SaaS-Vertrag, AVV, TOMs, Unterauftragnehmerliste und Löschkonzept freigegeben/unterzeichnet.", "Rechtsgrundlage und landesspezifisches Schulrecht dokumentiert.", "DSFA-Schwellwertanalyse beziehungsweise vollständige DSFA abgeschlossen.", "Informationspflichten gegenüber Schülern/Erziehungsberechtigten/Personal erfüllt.", "Nutzerkreis, Rollen und Kontoimport von autorisierter Stelle bestätigt.", "Verbot medizinischer Falldokumentation und zulässige Chat-/Dateiinhalte kommuniziert.", "Notfall-, Datenschutz- und Supportkontakte getestet.", "Export-, Austritts- und Kontolöschprozess festgelegt.", "Store-/App-Datenschutzinformation zugänglich.", "Freigabedatum und nächste Jahresprüfung terminiert."])]),
    ("Aktivierungsentscheidung", [("p", "[ ] Freigegeben  [ ] Freigegeben mit Auflagen  [ ] Nicht freigegeben\nAuflagen/Begründung: ______________________________________________________________________________\nAktivierungsdatum: __________________  Nächste Überprüfung: __________________")]),
], sign=True)
register("14_SSD_Manager_Schul_Onboarding_Freigabe.docx", "Onboarding-Freigabe", "Keine offene Registrierung; Schulprüfung dokumentieren", "Je Schule ausfüllen")

save("15_SSD_Manager_App_Store_Datenschutzangaben.docx", "App-Store-Datenschutzangaben", "Arbeitsblatt für Apple App Privacy und Google Play Data Safety", [
    ("Produkt und Kontozugang", [("p", "Bundle/Application ID: com.minutmate.ssdmanager. Öffentliche Distribution, aber kein frei zugängliches Konto und keine Selbstregistrierung. Schulen beantragen die Einrichtung über info@ssd-manager.minutmate.com; Konten werden nur durch autorisierte Stellen angelegt.")]),
    ("Erhobene Datentypen", [("table", ["Store-Kategorie", "SSD-Manager-Nutzung", "Zweck / Verknüpfung"], [
        ("Name / Benutzerkennung / E-Mail", "Ja", "Konto, Schule, Rolle; mit Nutzer verknüpft; Kernfunktion"),
        ("Nutzerinhalte", "Ja", "Ankündigungstexte, Fotos und Dateien; mit Konto/Schule verknüpft; Kernfunktion"),
        ("App-Aktivität", "Ja", "Dienstzuweisungen, Rollenaktionen, Sicherheits-/Auditereignisse; Kernfunktion/Sicherheit"),
        ("Geräte- oder andere Kennungen", "Ja", "Zufällige Installationskennung, Push-Token; Sicherheit und Benachrichtigungen"),
        ("Diagnose-/technische Daten", "Begrenzt", "Serverseitige Fehler-, Login- und Zustellprotokolle; Sicherheit/Betrieb"),
        ("Gesundheit", "Nein, nicht vorgesehen", "Keine Patientendaten, Diagnosen oder Behandlungsdokumentation zulässig"),
        ("Standort, Kontakte, Werbung, Tracking", "Nein", "Keine entsprechende App-Funktion oder Werbe-/Tracking-SDKs"),
    ], [2300, 1900, 5160])]),
    ("Weitergabe und Sicherheit", [("bullets", ["Railway verarbeitet App- und Datenbankdaten als Hostinganbieter.", "Firebase/Google verarbeitet Push-Token und technische Push-Metadaten.", "Apple/APNs verarbeitet iOS-Pushdaten; Google/FCM Android-Pushdaten.", "All-Inkl speichert ausschließlich clientseitig/authentifiziert verschlüsselte Backups.", "Produktive Übertragung erfolgt verschlüsselt; Konten können durch autorisierte Schulrollen gelöscht werden."])]),
    ("Vor Einreichung erneut prüfen", [("bullets", ["Apple- und Google-Fragebogen anhand der dann eingebauten SDK-Versionen beantworten.", "Datenschutz-URL und Support-URL öffentlich erreichbar machen.", "Account-Deletion-Erklärung: keine Selbstbedienung durch Minderjährige; schulischer Prozess mit Lehreraufsicht/Sani-Leitung.", "Keine Aussage „Daten werden nicht erhoben“, da Konto-, Inhalts- und Geräteinformationen verarbeitet werden.", "Screenshots und Store-Texte dürfen keine echten Schülerdaten zeigen."])]),
])
register("15_SSD_Manager_App_Store_Datenschutzangaben.docx", "Store-Datenschutzblatt", "Apple/Google-Fragebögen vorbereiten", "Bei Upload gegen Binärdatei prüfen")

save("00_SSD_Manager_Dokumentenregister.docx", "Dokumentenregister SSD Manager", "Übersicht, Übernahmebewertung und Freigabestatus", [
    ("Paketübersicht", [("table", ["Datei", "Zweck", "Status"], [(f, p, s) for f, _, p, s in documents], [3600, 3300, 2460])]),
    ("Übernahme aus StudyConnect", [("table", ["Übernehmbar", "Neu/angepasst für SSD Manager"], [
        ("MinutMate-Anbieterangaben, Anschrift, Telefon, USt-IdNr. und allgemeine Kontaktstruktur", "Produktkontakte info@ssd-manager.minutmate.com und support@ssd-manager.minutmate.com"),
        ("Grundaufbau von Impressum, SaaS-Vertrag, AVV, TOMs, Unterauftragnehmerliste und Vertraulichkeitsverpflichtung", "Leistungsbild Schulsanitätsdienst, Rollen, Datenkategorien und ausdrücklicher Ausschluss medizinischer Falldaten"),
        ("Verantwortlichkeitsmodell Schule = Verantwortlicher / MinutMate = Auftragsverarbeiter", "SSD-spezifische Diensthistorie, Ankündigungen, Anhänge, Accountlöschung und Anonymisierung"),
        ("Vorhandene DPA-Nachweise für Railway, Google/Firebase und All-Inkl als Ausgangsnachweise", "Tatsächliche SSD-Konto-/Projektzuordnung, Regionen, Unterauftragnehmer und aktuelle Fassungen belegen"),
        ("Unterschriften- und Prüfstatuskonzept", "30-Tage-Backuprotation, XChaCha20-Poly1305, eigener FTP-Bereich und erfolgreicher Restore-Test"),
    ], [4300, 5060])]),
    ("Nicht übernommen", [("bullets", ["StudyConnect-Lernhilfe, Lerngruppen, private Chats, Moderation und öffentliche Webfunktionen.", "StudyConnect-Anbieter Vercel, Resend und Firebase Storage, da sie im SSD-Manager-Produktivbetrieb nicht aktiv sind.", "StudyConnect-URLs, Supportadressen, Verschlüsselungsbehauptungen und alte Backup-/Restore-Abweichungen.", "Pilot- und Rechnungsvorlagen, bis Vertriebsmodell und Preis ausdrücklich festgelegt sind."])]),
    ("Rechts- und Freigabequellen", [("bullets", ["Art. 5, 12–14, 28, 30, 32–36 und 44 ff. DSGVO.", "§ 5 Digitale-Dienste-Gesetz für Anbieterinformationen.", "EDPB Guidelines 07/2020 zu Verantwortlichen und Auftragsverarbeitern.", "DSK/BfDI-Orientierungshilfe zu Online-Lernplattformen und Art.-35-Prüfung bei minderjährigen Betroffenen.", "Jeweiliges Landes-Schul- und Landesdatenschutzrecht der Vertragsschule."])]),
    ("Freigabereihenfolge", [("numbered", ["MinutMate-Anbieterangaben, Steuerstatus, Leistungs-/Preismodell und URLs bestätigen.", "DPA-Nachweise und Unterauftragnehmerlisten projektbezogen aktualisieren.", "Anwaltliche Prüfung von Impressum, Datenschutz, Nutzungsbedingungen, SaaS-Vertrag und AVV.", "TOMs/Incident/Restore durch technische Prüfung bestätigen.", "Schule/Schulträger und Datenschutzbeauftragten einbinden; DSFA-Entscheidung dokumentieren.", "Finale Versionen datieren, unterschreiben, öffentlich bereitstellen und jährlich prüfen."])]),
], note=True)

print(f"Created {len(list(OUT.glob('*.docx')))} documents in {OUT}")
