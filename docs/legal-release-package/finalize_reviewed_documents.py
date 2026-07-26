"""Apply narrowly scoped final details to the manually reviewed DOCX files.

This script intentionally does not rebuild the legal package. It edits only
known paragraphs and cells so manual changes made in Word remain authoritative.
"""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
WORD = ROOT / "word"


REPLACEMENTS = {
    "00_SSD_Manager_Dokumentenregister.docx": {
        "Pilot- und Rechnungsvorlagen, bis Vertriebsmodell und Preis ausdrücklich festgelegt sind.":
        "Standardisierte Preis-, Laufzeit- und Pilotvorgaben; diese werden mit jeder Schule beziehungsweise jedem Schulträger individuell vereinbart.",
        "MinutMate-Anbieterangaben, Steuerstatus, Leistungs-/Preismodell und URLs bestätigen.":
        "MinutMate-Anbieterangaben, Steuerstatus und öffentliche Produkt-URLs sind dokumentiert; kommerzielle Konditionen werden je Schule individuell vereinbart.",
        "Preis/Laufzeit und Recht prüfen": "Individuelle Konditionen ergänzen und Recht prüfen",
    },
    "01_SSD_Manager_Impressum.docx": {
        "Dieses Impressum gilt für die SSD-Manager-App, ihre produktbezogenen Webseiten und Subdomains. Vor Veröffentlichung sind die endgültigen URLs für Datenschutzerklärung, Nutzungsbedingungen und Support einzutragen.":
        "Dieses Impressum gilt für die SSD-Manager-App, die Website https://ssd-manager.minutmate.com und die zugehörigen produktbezogenen Seiten. Datenschutzerklärung: https://ssd-manager.minutmate.com/datenschutz/ · Nutzungsbedingungen: https://ssd-manager.minutmate.com/nutzungsbedingungen/ · Support: https://ssd-manager.minutmate.com/support/.",
    },
    "02_SSD_Manager_Datenschutzerklaerung.docx": {
        "Die wesentlichen Sicherheitsmaßnahmen sind in den TOMs beschrieben. Diese Erklärung ist bei Funktions-, Anbieter-, Rechts- oder Speicheränderungen zu aktualisieren. Vor Veröffentlichung sind Versionsdatum, Store-URLs und die verantwortliche Schule beziehungsweise ein schulbezogener Hinweis einzusetzen.":
        "Die wesentlichen Sicherheitsmaßnahmen sind in den TOMs beschrieben. Diese Erklärung ist bei Funktions-, Anbieter-, Rechts- oder Speicheränderungen zu aktualisieren. Die aktuelle öffentliche Fassung ist unter https://ssd-manager.minutmate.com/datenschutz/ erreichbar. Die jeweils verantwortliche Schule oder der Schulträger teilt den Nutzenden die eigenen Kontaktdaten und die konkrete schulrechtliche Rechtsgrundlage bei Bereitstellung der Schulumgebung mit.",
    },
    "04_SSD_Manager_SaaS_Vertrag_Schule.docx": {
        "Laufzeit: [ ] Schulsemester  [ ] Schuljahr  [ ] abweichend: __________________. Ordentliche Kündigungsfrist: 14 Tage zum Laufzeitende, soweit individuell nichts anderes vereinbart ist. Außerordentliche Kündigung aus wichtigem Grund bleibt möglich.":
        "Vertragsbeginn, Laufzeit, Verlängerung und ordentliche Kündigungsfristen werden mit jeder Schule beziehungsweise jedem Schulträger individuell im Angebot oder in einer Vertragsanlage vereinbart. Außerordentliche Kündigung aus wichtigem Grund bleibt möglich.",
        "________________ EUR": "Gemäß individuellem Angebot",
        "________________": "Gemäß individuellem Angebot",
        "14 Tage ab Rechnungseingang": "Gemäß individuellem Angebot",
    },
    "07_SSD_Manager_Unterauftragnehmerliste.docx": {
        "MySQL-Datenbankbetrieb und Hosting-Infrastruktur": "Ausschließlich verschlüsselte Offsite-Datenbanksicherungen in getrenntem FTP-Bereich",
        "Stamm-, Konto-, Rollen-, Schul-, Kommunikations-, Nutzungs- und Protokolldaten": "Verschlüsseltes Backup-Archiv; Anbieter erhält keinen Klartextinhalt",
        "Dateien, Dateimetadaten, technische Zugriffs- und Sicherheitsdaten": "Geräte-/Push-Token und technische Push-Metadaten",
        "europe-west3, Frankfurt": "EU-Verarbeitung soweit konfigurierbar; Push-Zustellung kann globale Infrastruktur einbeziehen",
        "Aktiv; Restore-Test in EU zu verlagern": "Aktiv; API und MySQL in EU West (Amsterdam)",
    },
    "11_SSD_Manager_Betroffenenanfragen_Prozess.docx": {
        "Produkt: support@ssd-manager.minutmate.com · Datenschutz MinutMate: datenschutz@minutmate.com · Schule/DSB: __________________. Aktenzeichen, Eingang, Identitätsprüfung, Verantwortlicher, Maßnahmen, Freigabe, Antwortdatum und Löschdatum sind zu dokumentieren.":
        "Produkt: support@ssd-manager.minutmate.com · Datenschutz MinutMate: datenschutz@minutmate.com · Schule/DSB: wird je Schulumgebung im Onboarding dokumentiert. Aktenzeichen, Eingang, Identitätsprüfung, Verantwortlicher, Maßnahmen, Freigabe, Antwortdatum und Löschdatum sind zu dokumentieren.",
    },
    "12_SSD_Manager_Datenschutzverletzung_Prozess.docx": {
        "Interner/technischer Meldeweg: support@ssd-manager.minutmate.com und verwaltung@minutmate.com. Datenschutzkontakt: datenschutz@minutmate.com. Notfallkontakt der Schule/Schulträger: __________________. MinutMate informiert die verantwortliche Stelle unverzüglich nach Bekanntwerden eines Vorfalls im Auftragsbereich.":
        "Interner/technischer Meldeweg: support@ssd-manager.minutmate.com und verwaltung@minutmate.com. Datenschutzkontakt: datenschutz@minutmate.com. Der Notfallkontakt der Schule beziehungsweise des Schulträgers wird je Schulumgebung im Onboarding dokumentiert. MinutMate informiert die verantwortliche Stelle unverzüglich nach Bekanntwerden eines Vorfalls im Auftragsbereich.",
    },
    "15_SSD_Manager_App_Store_Datenschutzangaben.docx": {
        "Datenschutz-URL https://ssd-manager.minutmate.com/datenschutz/ und Support-URL https://ssd-manager.minutmate.com/support/ vor Einreichung öffentlich erreichbar machen und in beiden Stores eintragen.":
        "Datenschutz-URL https://ssd-manager.minutmate.com/datenschutz/, Support-URL https://ssd-manager.minutmate.com/support/ und Erläuterung zur Kontolöschung https://ssd-manager.minutmate.com/konto-loeschen/ vor Einreichung öffentlich erreichbar machen und in beiden Stores eintragen.",
    },
}

COMPACT_DOCUMENTS = {
    "01_SSD_Manager_Impressum.docx",
    "11_SSD_Manager_Betroffenenanfragen_Prozess.docx",
    "12_SSD_Manager_Datenschutzverletzung_Prozess.docx",
    "13_SSD_Manager_Vertraulichkeitsverpflichtung.docx",
}


def paragraphs_in(container):
    yield from container.paragraphs
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from paragraphs_in(cell)


def replace_paragraph_text(paragraph, new_text):
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text


def main():
    for filename, replacements in REPLACEMENTS.items():
        path = WORD / filename
        document = Document(path)
        found = set()
        for paragraph in paragraphs_in(document):
            old = paragraph.text.strip()
            if old in replacements:
                replace_paragraph_text(paragraph, replacements[old])
                found.add(old)
        present_new = {old for old, new in replacements.items() if any(p.text.strip() == new for p in paragraphs_in(document))}
        missing = set(replacements) - found - present_new
        if missing:
            raise RuntimeError(f"Expected text not found in {filename}: {sorted(missing)}")
        if filename in COMPACT_DOCUMENTS:
            for section in document.sections:
                section.top_margin = Inches(0.78)
                section.bottom_margin = Inches(0.72)
            normal = document.styles["Normal"]
            normal.paragraph_format.space_after = Pt(4)
            normal.paragraph_format.line_spacing = 1.04
        if filename == "13_SSD_Manager_Vertraulichkeitsverpflichtung.docx":
            for marker in document._element.xpath(".//w:lastRenderedPageBreak"):
                marker.getparent().remove(marker)
            for section in document.sections:
                section.page_width = Inches(8.27)
                section.page_height = Inches(11.69)
                section.top_margin = Inches(0.42)
                section.bottom_margin = Inches(0.36)
                section.header_distance = Inches(0.2)
                section.footer_distance = Inches(0.2)
            document.styles["Normal"].font.size = Pt(10)
            for style_name in ("Heading 1", "Heading 2", "Heading 3"):
                style = document.styles[style_name]
                style.paragraph_format.space_before = Pt(7)
                style.paragraph_format.space_after = Pt(3)
            for paragraph in document.paragraphs:
                if not paragraph.text.strip():
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 0.1
                    if paragraph.runs:
                        paragraph.runs[0].font.size = Pt(1)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_after = Pt(1)
                            paragraph.paragraph_format.line_spacing = 1.0
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
        document.save(path)
        print(f"Updated {filename}: {len(found)} targeted fields")


if __name__ == "__main__":
    main()
